# -*- coding: utf-8 -*-
"""紹介資料を Word の形にする（dist/センゴク盤_紹介資料.docx）

   中身は src/data/shiryo.js ひとつから取る。tools/shiryo-json.cjs が
   その表を JSON に落としてくれるので、ここではそれを読んで組むだけである。

   紙（PDF）と同じ順、同じ字面にする。見た目だけが Word の作法に従う。
"""
import json, os, sys
from docx import Document
from docx.shared import Pt, Mm, RGBColor, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

ROOT = os.path.dirname(os.path.dirname(os.path.dirname(os.path.abspath(__file__))))
DATA = json.load(open(os.path.join(ROOT, 'build', 'shiryo.json'), encoding='utf-8'))
DIST = os.path.join(ROOT, 'dist')
絵蔵 = os.path.join(DIST, 'tebiki', '大')
図蔵 = os.path.join(DIST, 'tebiki', 'zu')

墨 = RGBColor(0x26, 0x26, 0x2A)
薄 = RGBColor(0x6E, 0x6A, 0x62)
灰 = RGBColor(0x8A, 0x84, 0x78)
明朝 = "Hiragino Mincho ProN"
ゴシック = "Hiragino Sans"

doc = Document()

def 和字(run, 名=明朝):
    """欧文と和文の両方に書体を効かせる（Word は別々に持つため）"""
    run.font.name = 名
    run._element.rPr.rFonts.set(qn('w:eastAsia'), 名)

for s in doc.sections:
    s.page_width, s.page_height = Mm(210), Mm(297)
    s.top_margin, s.bottom_margin = Mm(20), Mm(18)
    s.left_margin, s.right_margin = Mm(18), Mm(18)

n = doc.styles['Normal']
n.font.size = Pt(10.5); n.font.color.rgb = 墨
n.font.name = 明朝; n.element.rPr.rFonts.set(qn('w:eastAsia'), 明朝)
n.paragraph_format.line_spacing = 1.6
n.paragraph_format.space_after = Pt(5)

def 段(文, 大きさ=10.5, 色=墨, 書体=明朝, 前=0, 後=5, 揃=None, 太=False, 間=None):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(前)
    p.paragraph_format.space_after = Pt(後)
    if 揃 is not None: p.alignment = 揃
    r = p.add_run(文); r.font.size = Pt(大きさ); r.font.color.rgb = 色; r.bold = 太
    和字(r, 書体)
    if 間 is not None:
        rPr = r._element.get_or_add_rPr()
        sp = OxmlElement('w:spacing'); sp.set(qn('w:val'), str(int(間 * 20))); rPr.append(sp)
    return p

def 罫(p, 位置='bottom', 太さ=6, 色='26262A'):
    pPr = p._p.get_or_add_pPr()
    bd = pPr.find(qn('w:pBdr'))
    if bd is None:
        bd = OxmlElement('w:pBdr'); pPr.append(bd)
    e = OxmlElement('w:' + 位置)
    e.set(qn('w:val'), 'single'); e.set(qn('w:sz'), str(太さ))
    e.set(qn('w:space'), '4'); e.set(qn('w:color'), 色)
    bd.append(e)

# ------------------------------------------------------------------ 表紙
題 = os.path.join(ROOT, 'build', 'logo-title.png')
if os.path.exists(題):
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(150)
    p.add_run().add_picture(題, width=Mm(120))
段(DATA['資料名'], 14, 墨, 明朝, 前=24, 後=6, 揃=WD_ALIGN_PARAGRAPH.CENTER, 間=3)
段(DATA['添え書き'], 10, 薄, 明朝, 後=0, 揃=WD_ALIGN_PARAGRAPH.CENTER, 間=1)

段('目次', 9, 灰, ゴシック, 前=40, 後=6, 間=2)
for i, c in enumerate(DATA['資料']):
    p = doc.add_paragraph(); p.paragraph_format.space_after = Pt(2)
    r = p.add_run('%2d　' % (i + 1)); r.font.size = Pt(9); r.font.color.rgb = 灰; 和字(r, ゴシック)
    r = p.add_run(c['題']); r.font.size = Pt(11); 和字(r)
    if c.get('副'):
        r = p.add_run('　' + c['副']); r.font.size = Pt(9); r.font.color.rgb = 灰; 和字(r)

# ------------------------------------------------------------------ 本文
def 表を組む(t):
    行 = t.get('行') or []
    if not 行: return
    列 = len(行[0])
    頭あり = any(t.get('頭') or [])
    tb = doc.add_table(rows=0, cols=列)
    tb.style = 'Table Grid'
    tb.alignment = WD_TABLE_ALIGNMENT.LEFT
    # 罫を消して、下線だけ自前で引く
    for b in tb._tbl.iter(qn('w:tblBorders')):
        for e in list(b): b.remove(e)
        for 名 in ('top', 'left', 'bottom', 'right', 'insideH', 'insideV'):
            x = OxmlElement('w:' + 名); x.set(qn('w:val'), 'nil'); b.append(x)
    def 下線(cell, 太さ=4, 色='D8D2C4'):
        tcPr = cell._tc.get_or_add_tcPr()
        bd = OxmlElement('w:tcBorders')
        e = OxmlElement('w:bottom')
        e.set(qn('w:val'), 'single'); e.set(qn('w:sz'), str(太さ)); e.set(qn('w:color'), 色)
        bd.append(e); tcPr.append(bd)
    if 頭あり:
        cells = tb.add_row().cells
        for j, h in enumerate(t['頭']):
            cells[j].text = ''
            p = cells[j].paragraphs[0]; p.paragraph_format.space_after = Pt(2)
            r = p.add_run(h); r.font.size = Pt(8.5); r.font.color.rgb = 薄; r.bold = True
            和字(r, ゴシック)
            下線(cells[j], 8, '26262A')
    for row in 行:
        cells = tb.add_row().cells
        for j, v in enumerate(row):
            cells[j].text = ''
            p = cells[j].paragraphs[0]
            p.paragraph_format.space_before = Pt(2); p.paragraph_format.space_after = Pt(2)
            p.paragraph_format.line_spacing = 1.45
            r = p.add_run(v); r.font.size = Pt(9.5); 和字(r)
            if j == 0: r.bold = True
            else: r.font.color.rgb = RGBColor(0x4A, 0x46, 0x40)
            下線(cells[j])
    doc.add_paragraph().paragraph_format.space_after = Pt(3)

def 数を組む(k):
    tb = doc.add_table(rows=0, cols=2)
    tb.alignment = WD_TABLE_ALIGNMENT.LEFT
    for i in range(0, len(k), 2):
        cells = tb.add_row().cells
        for j, 項 in enumerate(k[i:i + 2]):
            名, 値, 添 = (項 + ['', '', ''])[:3]
            c = cells[j]; c.text = ''
            p = c.paragraphs[0]; p.paragraph_format.space_after = Pt(0)
            r = p.add_run(名); r.font.size = Pt(8.5); r.font.color.rgb = 灰; 和字(r, ゴシック)
            p2 = c.add_paragraph(); p2.paragraph_format.space_after = Pt(0)
            r = p2.add_run(値); r.font.size = Pt(15); 和字(r)
            if 添:
                p3 = c.add_paragraph(); p3.paragraph_format.space_after = Pt(4)
                r = p3.add_run(添); r.font.size = Pt(8.5); r.font.color.rgb = 薄; 和字(r)
    doc.add_paragraph().paragraph_format.space_after = Pt(3)

def 図を組む(z):
    for key, 名, 説 in z:
        f = os.path.join(図蔵, key + '.jpg')
        tb = doc.add_table(rows=1, cols=2)
        tb.alignment = WD_TABLE_ALIGNMENT.LEFT
        c0, c1 = tb.rows[0].cells
        c0.width, c1.width = Mm(28), Mm(146)
        c0.text = ''
        if os.path.exists(f):
            c0.paragraphs[0].add_run().add_picture(f, width=Mm(26))
        c1.text = ''
        p = c1.paragraphs[0]; p.paragraph_format.space_after = Pt(1)
        r = p.add_run(名); r.font.size = Pt(11); r.bold = True; 和字(r)
        p2 = c1.add_paragraph(); p2.paragraph_format.space_after = Pt(3)
        p2.paragraph_format.line_spacing = 1.5
        r = p2.add_run(説); r.font.size = Pt(9.5); r.font.color.rgb = RGBColor(0x4A, 0x46, 0x40)
        和字(r)

for i, c in enumerate(DATA['資料']):
    doc.add_page_break()
    p = doc.add_paragraph(); p.paragraph_format.space_after = Pt(10)
    r = p.add_run('%d　' % (i + 1)); r.font.size = Pt(10); r.font.color.rgb = 灰; 和字(r, ゴシック)
    r = p.add_run(c['題']); r.font.size = Pt(17); 和字(r)
    if c.get('副'):
        r = p.add_run('　　' + c['副']); r.font.size = Pt(10); r.font.color.rgb = 灰; 和字(r, ゴシック)
    罫(p, 太さ=12)

    for s in c['節']:
        if s.get('見出し'):
            段(s['見出し'], 11.5, RGBColor(0x4A, 0x46, 0x40), ゴシック, 前=12, 後=3, 太=True, 間=1.2)
        if s.get('絵'):
            f = os.path.join(絵蔵, s['絵'] + '.jpg')
            if os.path.exists(f):
                p = doc.add_paragraph(); p.paragraph_format.space_after = Pt(2)
                p.add_run().add_picture(f, width=Mm(168))
                if s.get('絵の説'):
                    段(s['絵の説'], 8.5, 薄, 明朝, 後=8)
        for t in s.get('文') or []:
            p = 段(t)
            p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        if s.get('表'): 表を組む(s['表'])
        if s.get('数'): 数を組む(s['数'])
        if s.get('図'): 図を組む(s['図'])
        for t in s.get('箇条') or []:
            p = doc.add_paragraph(style='List Bullet')
            p.paragraph_format.space_after = Pt(2)
            p.paragraph_format.line_spacing = 1.5
            r = p.add_run(t); r.font.size = Pt(10); 和字(r)

out = os.path.join(DIST, '%s_%s.docx' % (DATA['題名'], DATA['資料名']))
doc.save(out)
print('dist/%s    %d KB' % (os.path.basename(out), os.path.getsize(out) // 1024))
